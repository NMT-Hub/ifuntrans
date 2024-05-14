import argparse
import asyncio
import re

from docx import Document
from ifuntrans.translate import translate
from typing import List
from ifuntrans.tm import create_tm_from_excel


async def main():
    parser = argparse.ArgumentParser(description="Translate rtf file")
    parser.add_argument("file", help="The excel file to translate")
    parser.add_argument("-f", "--from-lang", help="The source language", default="en", type=str)
    parser.add_argument("-t", "--to-lang", help="The target language", default="zh", type=str)
    parser.add_argument("--output", help="The output file prefix", default=None, type=str)
    parser.add_argument("-tm", "--translate-memory-file", help="The translate memory file", default=None, type=str)
    parser.add_argument("-i", "--instructions", help="The instructions prompt", default="", type=str)
    args = parser.parse_args()

    document = Document(args.file)
    if args.translate_memory_file is not None:
        tm = await create_tm_from_excel(args.translate_memory_file)
    else:
        tm = None
    
    # TODO Semple merge paragraphs, only CJK works. Use mBert to merge paragraphs more accurately
    # merge the neighboring paragraphs if they have the same style and 
    # final character in the first paragraph is CJK and the first character in the second paragraph is CJK
    for i in range(len(document.paragraphs) - 1):
        if document.paragraphs[i].text == "" or document.paragraphs[i + 1].text == "":
            continue
        if document.paragraphs[i].style == document.paragraphs[i + 1].style:
            if re.match(r"[\u4e00-\u9fa5]", document.paragraphs[i].text[-1]) and re.match(r"[\u4e00-\u9fa5]", document.paragraphs[i + 1].text[0]):
                document.paragraphs[i].text += " " + document.paragraphs[i + 1].text
                document.paragraphs[i + 1].text = ""
    
    texts = []
    for paragraph in document.paragraphs:
        texts.append(paragraph.text)

    translation: List[str] = await translate(
            texts,
            from_lang=args.from_lang,
            to_lang=args.to_lang,
            tm=tm,
            instructions=args.instructions,
        )

    for i, paragraph in enumerate(document.paragraphs):
        paragraph.text = translation[i]
        
    # translate table text
    table_texts = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                table_texts.append(cell.text)
                
    translation: List[str] = await translate(
            table_texts,
            from_lang=args.from_lang,
            to_lang=args.to_lang,
            tm=tm,
            instructions=args.instructions,
        )
    
    for i, table in enumerate(document.tables):
        for j, row in enumerate(table.rows):
            for k, cell in enumerate(row.cells):
                cell.text = translation[i * len(table.rows) + j * len(row.cells) + k]

    if args.output is None:
        args.output = args.file

    # save
    document.save(args.output)


if __name__ == "__main__":
    asyncio.run(main())
